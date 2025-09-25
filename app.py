from flask import Flask, render_template, request, jsonify
import re
from datetime import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.naive_bayes import MultinomialNB
import numpy as np
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.chunk import ne_chunk
from nltk.tag import pos_tag
from docx import Document
import os

# Try to import pdfplumber with fallback
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    print("Warning: pdfplumber not available. PDF processing will be limited.")
    PDF_AVAILABLE = False

app = Flask(__name__)

# Try to load spaCy model with fallback
try:
    import spacy
    nlp = spacy.load('en_core_web_sm')
    print("SpaCy model loaded successfully")
except Exception as e:
    print(f"Warning: Could not load spaCy model: {e}")
    print("Creating fallback NLP processor")
    nlp = None

# Initialize NLTK components
try:
    # Try to use the data first
    import nltk.data
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('vader_lexicon')
    nltk.data.find('averaged_perceptron_tagger')
    nltk.data.find('maxent_ne_chunker')
    nltk.data.find('words')
except LookupError:
    # Download if not found
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('punkt_tab', quiet=True)  # Download new punkt_tab
        nltk.download('stopwords', quiet=True)
        nltk.download('vader_lexicon', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
        nltk.download('maxent_ne_chunker', quiet=True)
        nltk.download('words', quiet=True)
    except Exception as e:
        print(f"Warning: Could not download NLTK data: {e}")

# Initialize sentiment analyzer with error handling
try:
    sia = SentimentIntensityAnalyzer()
    stop_words = set(stopwords.words('english'))
except Exception as e:
    print(f"Warning: NLTK initialization error: {e}")
    sia = None
    stop_words = set()

# Fallback NLP processor when spaCy is not available
class FallbackNLP:
    def __init__(self):
        self.entities = []
    
    def __call__(self, text):
        return FallbackDoc(text)

class FallbackDoc:
    def __init__(self, text):
        self.text = text
        self.ents = self._extract_entities(text)
    
    def _extract_entities(self, text):
        """Extract basic entities without spaCy"""
        entities = []
        
        # Simple name extraction - look for capitalized words at the beginning
        lines = text.split('\n')[:10]
        for line in lines:
            words = line.strip().split()
            if len(words) >= 2:
                # Check if first two words are capitalized (likely names)
                if all(word[0].isupper() and word.isalpha() for word in words[:2]):
                    entity = FallbackEntity(' '.join(words[:2]), 'PERSON')
                    entities.append(entity)
        
        # Simple location detection - common city patterns
        import re
        cities = ['mumbai', 'delhi', 'bangalore', 'chennai', 'hyderabad', 'pune', 'kolkata', 'ahmedabad', 'puducherry']
        for city in cities:
            if city in text.lower():
                entity = FallbackEntity(city.title(), 'GPE')
                entities.append(entity)
        
        # Date patterns
        date_pattern = r'\b\d{4}\b'
        dates = re.findall(date_pattern, text)
        for date in dates[:3]:  # Limit to first 3 dates
            entity = FallbackEntity(date, 'DATE')
            entities.append(entity)
        
        return entities

class FallbackEntity:
    def __init__(self, text, label):
        self.text = text
        self.label_ = label

# Use fallback if spaCy is not available
if nlp is None:
    nlp = FallbackNLP()
    print("Using fallback NLP processor")
    # Fallback: use empty sentiment analyzer
    class FallbackSentimentAnalyzer:
        def polarity_scores(self, text):
            return {'pos': 0.0, 'neg': 0.0, 'neu': 1.0, 'compound': 0.0}
    
    sia = FallbackSentimentAnalyzer()
    stop_words = set()

IGNORE_LIST = [
    "ABOUT ME", "OBJECTIVE", "SUMMARY", "RESUME", "SKILLS", "EDUCATION", "PROJECTS",
    "EXPERIENCE", "CERTIFICATIONS", "INTERESTS", "LANGUAGES", "CONTACT", "PROFILE",
    "PYTHON", "JAVA", "JAVASCRIPT", "HTML", "CSS", "C++", "SQL", "DATA SCIENCE",
    "MSWORD", "POWERPOINT", "TALLY", "MYSQL"
]

# Enhanced training data for job role prediction with more comprehensive patterns
JOB_TRAINING_DATA = {
    "Data Scientist": [
        "python machine learning data analysis pandas numpy scikit-learn statistics modeling predictive analytics",
        "sql python data mining statistical analysis machine learning algorithms deep learning neural networks",
        "data visualization tableau powerbi python r statistics predictive modeling analytics data insights",
        "machine learning tensorflow pytorch neural networks data science python statistics mathematical modeling",
        "big data hadoop spark python data analysis machine learning statistical modeling data processing",
        "data analytics business intelligence statistics machine learning python sql data warehousing",
        "predictive modeling regression classification clustering statistical analysis data mining algorithms",
        "artificial intelligence deep learning computer vision natural language processing data science",
        "statistical computing r python data analysis hypothesis testing experimental design analytics"
    ],
    "Software Engineer": [
        "java spring boot microservices rest api backend development software engineering object oriented programming",
        "python django flask web development backend apis database programming software architecture",
        "javascript react nodejs full stack development web applications frontend backend software development",
        "c++ software development algorithms data structures object oriented programming system programming",
        "software engineering design patterns testing debugging version control git agile development",
        "full stack development web applications database design software architecture programming languages",
        "backend development api design microservices software engineering database integration",
        "software development lifecycle agile methodologies version control testing deployment automation",
        "programming languages software engineering web development mobile applications system design"
    ],
    "Frontend Developer": [
        "html css javascript react angular vue frontend development responsive design user interface",
        "javascript typescript react redux frontend web development user interface user experience",
        "css html sass bootstrap responsive design frontend development user experience mobile first",
        "react angular javascript frontend development single page applications ui ux design",
        "web development frontend javascript css html responsive design mobile first progressive web apps",
        "user interface design javascript frontend frameworks react vue angular responsive web design",
        "frontend development html5 css3 javascript dom manipulation cross browser compatibility",
        "responsive web design mobile first css preprocessors javascript frameworks ui components",
        "web accessibility frontend optimization performance tuning user experience design"
    ],
    "Backend Developer": [
        "java spring boot rest api microservices backend development database design server programming",
        "python django flask rest api backend development database postgresql mysql server architecture",
        "nodejs express mongodb rest api backend development server side programming database design",
        "database design sql postgresql mysql backend development api design data modeling",
        "microservices architecture backend development rest api database design system architecture",
        "server side development api design database optimization backend programming system design",
        "rest api graphql backend development microservices database integration server architecture",
        "backend programming server development database design api development system optimization",
        "cloud computing backend development aws azure database management server infrastructure"
    ],
    "DevOps Engineer": [
        "docker kubernetes aws cloud infrastructure deployment automation ci cd pipeline management",
        "linux bash scripting automation infrastructure docker containers orchestration system administration",
        "aws azure cloud infrastructure terraform ansible devops automation deployment continuous integration",
        "jenkins ci cd pipeline automation docker kubernetes infrastructure as code deployment strategies",
        "monitoring logging infrastructure automation cloud platforms devops practices site reliability",
        "containerization docker kubernetes microservices orchestration cloud native infrastructure",
        "infrastructure as code terraform ansible configuration management cloud automation devops",
        "continuous integration continuous deployment jenkins gitlab ci automation testing deployment",
        "cloud computing aws azure gcp infrastructure monitoring system administration devops practices"
    ],
    "Mobile Developer": [
        "android java kotlin mobile development android studio mobile applications native development",
        "ios swift xcode mobile development iphone ipad applications app store native ios",
        "react native javascript mobile development cross platform ios android hybrid applications",
        "flutter dart mobile development cross platform android ios applications mobile ui",
        "mobile development android ios react native flutter cross platform native hybrid apps",
        "native mobile development android kotlin ios swift mobile applications user interface",
        "cross platform development react native flutter xamarin mobile applications hybrid development",
        "mobile app development android ios native hybrid progressive web applications mobile ui",
        "mobile development frameworks android studio xcode react native flutter mobile programming"
    ],
    "Data Engineer": [
        "data engineering etl pipeline apache spark hadoop big data processing data warehousing",
        "python sql data pipeline etl apache airflow data processing big data engineering",
        "apache spark hadoop kafka data streaming real time data processing big data",
        "data warehousing sql python etl processes data modeling data architecture",
        "big data processing apache spark hadoop data pipeline kafka streaming data engineering",
        "cloud data engineering aws azure data pipeline data lake data warehousing",
        "data architecture data modeling etl pipeline data integration big data processing",
        "streaming data processing kafka apache storm real time analytics data engineering",
        "data pipeline automation python sql apache airflow data processing data integration"
    ],
    "Machine Learning Engineer": [
        "machine learning tensorflow pytorch deep learning neural networks model deployment production",
        "python machine learning scikit-learn deep learning tensorflow model training deployment",
        "deep learning tensorflow pytorch computer vision natural language processing neural networks",
        "machine learning operations mlops model deployment model monitoring production systems",
        "artificial intelligence machine learning deep learning model development production deployment",
        "neural networks deep learning tensorflow pytorch computer vision nlp machine learning",
        "model deployment kubernetes docker machine learning production systems mlops automation",
        "machine learning algorithms supervised unsupervised learning model optimization production deployment",
        "deep learning frameworks tensorflow pytorch keras neural networks model training deployment"
    ],
    "Product Manager": [
        "product management roadmap strategy stakeholder management agile scrum product development",
        "product strategy market research user experience product development lifecycle business analysis",
        "agile scrum product owner requirements gathering stakeholder communication product roadmap",
        "market analysis competitive research product strategy business analysis product planning",
        "product roadmap feature prioritization user stories product development agile methodologies",
        "product management strategy market research competitive analysis stakeholder management",
        "user research product strategy business requirements agile product management",
        "product development lifecycle market analysis user experience business strategy",
        "product strategy business analysis market research product roadmap stakeholder management"
    ],
    "UI/UX Designer": [
        "user experience design user interface design wireframing prototyping design thinking",
        "ui ux design figma sketch adobe xd user interface design user experience",
        "user research usability testing user experience design interface design prototyping",
        "design systems ui components user interface design user experience prototyping",
        "wireframing prototyping user experience design user interface design design thinking",
        "user centered design usability testing user research interface design user experience",
        "visual design user interface ui ux design graphic design web design",
        "interaction design user experience design prototyping wireframing user interface",
        "design thinking user research user experience design interface design usability"
    ],
    "QA Engineer": [
        "quality assurance testing automation selenium test automation manual testing",
        "software testing test automation selenium webdriver quality assurance bug tracking",
        "test automation framework selenium cucumber java python testing quality assurance",
        "manual testing automated testing test planning quality assurance software testing",
        "testing methodologies agile testing test automation quality assurance regression testing",
        "software quality assurance test planning test execution bug tracking defect management",
        "automation testing selenium cypress api testing performance testing quality assurance",
        "test driven development quality assurance software testing test automation frameworks",
        "quality assurance testing selenium automation testing manual testing test planning"
    ],
    "Digital Marketing Specialist": [
        "digital marketing seo sem social media marketing content marketing analytics google ads",
        "marketing strategy brand management social media campaigns market research digital advertising",
        "content marketing copywriting seo social media digital marketing campaigns email marketing",
        "marketing analytics google analytics facebook ads digital advertising campaigns ppc marketing",
        "brand marketing social media strategy content creation marketing campaigns digital marketing",
        "seo search engine optimization content marketing digital marketing social media marketing",
        "digital advertising google ads facebook ads ppc marketing social media marketing",
        "marketing automation email marketing crm digital marketing campaign management analytics",
        "social media marketing content creation brand management digital marketing strategy"
    ]
}

# Initialize and train the ML model
def create_job_prediction_model():
    texts = []
    labels = []
    
    for job_role, samples in JOB_TRAINING_DATA.items():
        for sample in samples:
            texts.append(sample)
            labels.append(job_role)
    
    vectorizer = TfidfVectorizer(max_features=1000, stop_words='english', ngram_range=(1, 2))
    X = vectorizer.fit_transform(texts)
    
    model = LogisticRegression(random_state=42, max_iter=1000)
    model.fit(X, labels)
    
    return vectorizer, model

# Create the model globally
vectorizer, job_prediction_model = create_job_prediction_model()

def predict_job_role(resume_text):
    """Predict job role using enhanced ML model based on resume content"""
    try:
        import re
        from collections import Counter
        
        if not resume_text or len(resume_text.strip()) < 10:
            return "Unable to determine", 0
        
        # Enhanced text preprocessing
        def preprocess_text(text):
            # Convert to lowercase and remove extra whitespace
            text = re.sub(r'\s+', ' ', text.lower().strip())
            
            # Remove common resume noise words but keep technical terms
            noise_words = [
                'resume', 'cv', 'curriculum vitae', 'objective', 'summary',
                'references available upon request', 'phone', 'email', 'address',
                'linkedin', 'github', 'portfolio', 'website', 'profile'
            ]
            
            for noise in noise_words:
                text = re.sub(r'\b' + re.escape(noise) + r'\b', ' ', text)
            
            # Extract and emphasize technical skills
            tech_pattern = r'\b(?:python|java|javascript|typescript|react|angular|vue|node|nodejs|django|flask|spring|tensorflow|pytorch|sql|mysql|postgresql|mongodb|aws|azure|docker|kubernetes|git|html|css|sass|bootstrap|api|rest|graphql|microservices|agile|scrum|machine learning|data science|analytics|tableau|powerbi|selenium|jenkins|ci\/cd|devops|linux|bash|android|ios|swift|kotlin|flutter|xamarin|c\+\+|ruby|php|scala|hadoop|spark|kafka|airflow|redis|elasticsearch|nginx|apache)\b'
            tech_matches = re.findall(tech_pattern, text, re.IGNORECASE)
            
            # Emphasize technical skills by adding them multiple times
            emphasized_tech = ' '.join(tech_matches * 2)
            return f"{text} {emphasized_tech}"
        
        processed_text = preprocess_text(resume_text)
        
        # Prepare enhanced training data
        training_texts = []
        training_labels = []
        
        for role, descriptions in JOB_TRAINING_DATA.items():
            for desc in descriptions:
                training_texts.append(preprocess_text(desc))
                training_labels.append(role)
        
        # Enhanced TF-IDF vectorization
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.linear_model import LogisticRegression
        
        vectorizer = TfidfVectorizer(
            max_features=1000,
            stop_words='english',
            ngram_range=(1, 3),  # Include unigrams, bigrams, and trigrams
            min_df=1,
            max_df=0.8,
            lowercase=True,
            analyzer='word'
        )
        
        # Fit and transform training data
        X_train = vectorizer.fit_transform(training_texts)
        
        # Train enhanced model with better parameters
        model = LogisticRegression(
            random_state=42,
            max_iter=1000,
            C=1.0,
            class_weight='balanced'
        )
        model.fit(X_train, training_labels)
        
        # Transform and predict
        X_test = vectorizer.transform([processed_text])
        prediction = model.predict(X_test)[0]
        confidence = model.predict_proba(X_test)[0]
        
        # Get base confidence score
        base_confidence = max(confidence) * 100
        
        # Enhanced confidence calculation with keyword matching
        def calculate_enhanced_confidence(text, predicted_role):
            # Define role-specific keywords with weights
            role_keywords = {
                "Data Scientist": {
                    'high': ['machine learning', 'data science', 'statistical', 'analytics', 'modeling', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn'],
                    'medium': ['python', 'r', 'sql', 'statistics', 'data', 'analysis', 'visualization', 'tableau', 'powerbi'],
                    'low': ['research', 'insights', 'patterns', 'trends', 'hypothesis']
                },
                "Software Engineer": {
                    'high': ['software engineer', 'software development', 'programming', 'algorithms', 'data structures'],
                    'medium': ['java', 'python', 'c++', 'javascript', 'spring', 'django', 'flask', 'api', 'backend', 'frontend'],
                    'low': ['coding', 'development', 'application', 'system', 'architecture']
                },
                "Frontend Developer": {
                    'high': ['frontend', 'ui', 'user interface', 'responsive design', 'web development'],
                    'medium': ['html', 'css', 'javascript', 'react', 'angular', 'vue', 'sass', 'bootstrap'],
                    'low': ['design', 'user experience', 'mobile', 'browser', 'dom']
                },
                "Backend Developer": {
                    'high': ['backend', 'server', 'api', 'microservices', 'database design'],
                    'medium': ['java', 'python', 'nodejs', 'spring', 'django', 'express', 'sql', 'mongodb'],
                    'low': ['server side', 'integration', 'architecture', 'performance', 'scalability']
                },
                "DevOps Engineer": {
                    'high': ['devops', 'ci/cd', 'infrastructure', 'deployment', 'automation'],
                    'medium': ['docker', 'kubernetes', 'aws', 'azure', 'jenkins', 'terraform', 'ansible'],
                    'low': ['cloud', 'monitoring', 'scripting', 'pipeline', 'orchestration']
                },
                "Mobile Developer": {
                    'high': ['mobile development', 'android', 'ios', 'mobile app'],
                    'medium': ['java', 'kotlin', 'swift', 'react native', 'flutter', 'xamarin'],
                    'low': ['app store', 'mobile', 'native', 'cross platform', 'hybrid']
                },
                "Machine Learning Engineer": {
                    'high': ['machine learning engineer', 'ml', 'deep learning', 'neural networks', 'model deployment'],
                    'medium': ['tensorflow', 'pytorch', 'scikit-learn', 'python', 'computer vision', 'nlp'],
                    'low': ['ai', 'artificial intelligence', 'training', 'optimization', 'feature engineering']
                },
                "Data Engineer": {
                    'high': ['data engineer', 'data pipeline', 'etl', 'big data'],
                    'medium': ['spark', 'hadoop', 'kafka', 'airflow', 'python', 'sql'],
                    'low': ['data processing', 'warehousing', 'streaming', 'batch processing']
                },
                "Product Manager": {
                    'high': ['product manager', 'product management', 'roadmap', 'stakeholder management'],
                    'medium': ['agile', 'scrum', 'product owner', 'market research', 'strategy'],
                    'low': ['requirements', 'business analysis', 'user stories', 'prioritization']
                },
                "UI/UX Designer": {
                    'high': ['ui designer', 'ux designer', 'user experience', 'user interface', 'design thinking'],
                    'medium': ['figma', 'sketch', 'adobe xd', 'wireframing', 'prototyping'],
                    'low': ['usability', 'user research', 'visual design', 'interaction design']
                },
                "QA Engineer": {
                    'high': ['qa engineer', 'quality assurance', 'test automation', 'testing'],
                    'medium': ['selenium', 'cypress', 'automated testing', 'manual testing'],
                    'low': ['bug tracking', 'test planning', 'regression testing', 'defect management']
                },
                "Digital Marketing Specialist": {
                    'high': ['digital marketing', 'marketing specialist', 'seo', 'social media marketing'],
                    'medium': ['google ads', 'facebook ads', 'content marketing', 'analytics'],
                    'low': ['brand management', 'campaigns', 'email marketing', 'ppc']
                }
            }
            
            # Calculate keyword bonus
            bonus = 0
            if predicted_role in role_keywords:
                keywords = role_keywords[predicted_role]
                text_lower = text.lower()
                
                for keyword in keywords.get('high', []):
                    if keyword in text_lower:
                        bonus += 12
                for keyword in keywords.get('medium', []):
                    if keyword in text_lower:
                        bonus += 6
                for keyword in keywords.get('low', []):
                    if keyword in text_lower:
                        bonus += 2
            
            # Apply bonus with diminishing returns
            enhanced_confidence = base_confidence + min(bonus, 25)
            return min(enhanced_confidence, 92)  # Cap at 92%
        
        final_confidence = calculate_enhanced_confidence(resume_text, prediction)
        
        return prediction, int(final_confidence)
        
    except Exception as e:
        print(f"Job prediction error: {e}")
        return "General Software Engineer", 75

def analyze_resume_content(resume_text):
    """Analyze resume content and extract key information"""
    text_lower = resume_text.lower()
    
    # Extract skills
    skills = []
    skill_keywords = [
        'python', 'java', 'javascript', 'react', 'angular', 'vue', 'nodejs', 'html', 'css',
        'sql', 'mysql', 'postgresql', 'mongodb', 'docker', 'kubernetes', 'aws', 'azure',
        'machine learning', 'data science', 'tensorflow', 'pytorch', 'pandas', 'numpy',
        'spring boot', 'django', 'flask', 'express', 'git', 'jenkins', 'ci/cd'
    ]
    
    for skill in skill_keywords:
        if skill in text_lower:
            skills.append(skill.title())
    
    # Extract education level with enhanced detection
    education_level = "Not specified"
    
    # Use more precise pattern matching with word boundaries to avoid false matches
    import re
    
    # Check for specific degree patterns (more comprehensive and precise)
    if re.search(r'\b(phd|ph\.d|doctorate|doctoral|ph\s+d|doctor\s+of\s+philosophy)\b', text_lower):
        education_level = "PhD/Doctorate"
    elif re.search(r'\b(master|masters|master\'?s|mtech|m\.tech|m\s+tech|msc|m\.sc|m\s+sc|ms\b|mba|m\.b\.a|m\s+b\s+a|post\s+graduate|postgraduate|pg\b)\b', text_lower):
        education_level = "Master's Degree"
    elif re.search(r'\b(bachelor|bachelors|bachelor\'?s|btech|b\.tech|b\s+tech|bsc|b\.sc|b\s+sc|be\b|b\.e|b\s+e|bca|b\.ca|b\s+ca|bcom|b\.com|b\s+com|undergraduate|graduate\s+degree|ug\b)\b', text_lower):
        education_level = "Bachelor's Degree"
    elif re.search(r'\b(diploma|certificate|polytechnic|advanced\s+diploma|professional\s+certificate)\b', text_lower):
        education_level = "Diploma/Certificate"
    elif re.search(r'\b(12th|intermediate|higher\s+secondary|\+2|hsc|h\.s\.c|senior\s+secondary)\b', text_lower):
        education_level = "Higher Secondary"
    elif re.search(r'\b(10th|matriculation|high\s+school|sslc|s\.s\.l\.c|secondary)\b', text_lower):
        education_level = "High School"
    
    # Debug education detection
    print(f"DEBUG - Detected education level: {education_level}")
    
    # Extract experience level
    experience_level = "Entry Level"
    if any(word in text_lower for word in ['senior', '5+ years', '5 years', '6 years', '7 years', '8 years', '9 years']):
        experience_level = "Senior Level"
    elif any(word in text_lower for word in ['3 years', '4 years', 'mid-level', 'intermediate']):
        experience_level = "Mid Level"
    elif any(word in text_lower for word in ['intern', 'fresher', 'graduate', 'entry']):
        experience_level = "Entry Level"
    
    return {
        'skills': skills[:10],  # Top 10 skills
        'education_level': education_level,
        'experience_level': experience_level
    }

def analyze_resume_with_nltk(resume_text):
    """Enhanced resume analysis using NLTK with error handling"""
    try:
        # Tokenize text with fallback
        try:
            sentences = sent_tokenize(resume_text)
            words = word_tokenize(resume_text.lower())
        except Exception as e:
            print(f"Tokenization error: {e}")
            # Fallback: simple split
            sentences = resume_text.split('.')
            words = resume_text.lower().split()
        
        # Remove stopwords and non-alphabetic tokens
        filtered_words = [word for word in words if word.isalpha() and word not in stop_words]
        
        # Perform sentiment analysis with error handling
        try:
            sentiment_scores = sia.polarity_scores(resume_text)
        except Exception as e:
            print(f"Sentiment analysis error: {e}")
            sentiment_scores = {'pos': 0.0, 'neg': 0.0, 'neu': 1.0, 'compound': 0.0}
        
        # Extract named entities using NLTK with error handling
        nltk_entities = []
        try:
            tokens = word_tokenize(resume_text)
            pos_tags = pos_tag(tokens)
            chunks = ne_chunk(pos_tags)
            
            # Extract entities from chunks
            for chunk in chunks:
                if hasattr(chunk, 'label'):
                    entity_text = ' '.join([token for token, pos in chunk.leaves()])
                    nltk_entities.append({
                        'text': entity_text,
                        'label': chunk.label()
                    })
        except Exception as e:
            print(f"Named entity recognition error: {e}")
            nltk_entities = []
        
        # Calculate text statistics
        word_count = len(words)
        sentence_count = len(sentences)
        avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0
        
        # Extract key phrases (simple approach using POS tags) with error handling
        key_phrases = []
        try:
            current_phrase = []
            pos_tags = pos_tag(word_tokenize(resume_text))
            
            for word, pos in pos_tags:
                if pos.startswith('NN') or pos.startswith('JJ'):  # Nouns and adjectives
                    current_phrase.append(word)
                else:
                    if len(current_phrase) > 1:
                        phrase = ' '.join(current_phrase)
                        if len(phrase) > 3:  # Only meaningful phrases
                            key_phrases.append(phrase)
                    current_phrase = []
        except Exception as e:
            print(f"Key phrase extraction error: {e}")
            # Fallback: extract common phrases
            key_phrases = []
        
        return {
            'sentiment': {
                'positive': sentiment_scores['pos'],
                'negative': sentiment_scores['neg'],
                'neutral': sentiment_scores['neu'],
                'compound': sentiment_scores['compound']
            },
            'entities': nltk_entities,
            'statistics': {
                'word_count': word_count,
                'sentence_count': sentence_count,
                'avg_sentence_length': round(avg_sentence_length, 2)
            },
            'key_phrases': key_phrases[:10],  # Top 10 key phrases
            'filtered_words': filtered_words[:50]  # Top 50 meaningful words
        }
    
    except Exception as e:
        print(f"NLTK analysis error: {e}")
        # Return safe fallback data
        return {
            'sentiment': {
                'positive': 0.0,
                'negative': 0.0,
                'neutral': 1.0,
                'compound': 0.0
            },
            'entities': [],
            'statistics': {
                'word_count': len(resume_text.split()),
                'sentence_count': len(resume_text.split('.')),
                'avg_sentence_length': 0
            },
            'key_phrases': [],
            'filtered_words': []
        }

def extract_text(file_storage):
    filename = file_storage.filename.lower()
    if filename.endswith('.pdf'):
        if PDF_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_storage) as pdf:
                    text = ''
                    for page in pdf.pages:
                        text += page.extract_text() or ''
                return text
            except Exception as e:
                print(f"Error extracting text from PDF with pdfplumber: {e}")
                return f"PDF file uploaded: {file_storage.filename}. Please try uploading the resume as a text file or Word document."
        else:
            return f"PDF processing not available. Please try uploading the resume as a text file or Word document."
    elif filename.endswith('.docx'):
        try:
            doc = Document(file_storage)
            text = ""
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    else:
        try:
            return file_storage.read().decode('utf-8')
        except UnicodeDecodeError:
            file_storage.seek(0)
            return file_storage.read().decode('latin-1')

def is_likely_name(text):
    """Enhanced name detection with intelligent pattern recognition"""
    if not text or not text.strip():
        return False
        
    words = text.strip().split()
    text_clean = text.strip()
    text_upper = text_clean.upper()
    text_lower = text_clean.lower()
    
    # Basic validation - length and word count
    if (len(words) < 1 or len(words) > 4 or
        text_upper in IGNORE_LIST or
        any(char.isdigit() for char in text_clean) or
        len(text_clean) > 50):  # Names shouldn't be too long
        return False
    
    # Company/Organization indicators (definitive non-names)
    company_indicators = {
        'company', 'corp', 'corporation', 'inc', 'incorporated', 'ltd', 'limited', 'llc',
        'solutions', 'systems', 'technologies', 'services', 'consulting', 'software',
        'tech', 'it', 'information', 'technology', 'pvt', 'private', 'public',
        'enterprise', 'enterprises', 'group', 'international', 'global', 'india',
        'industries', 'foundation', 'institute', 'organization', 'association'
    }
    
    # Check if any word in the text is a company indicator
    for word in words:
        if word.lower() in company_indicators:
            return False
    
    # Technical terms that are definitely not names
    technical_terms = {
        'api', 'ui', 'ux', 'css', 'html', 'sql', 'xml', 'json', 'rest', 'soap', 'http', 'https',
        'tcp', 'ip', 'dns', 'url', 'uri', 'cdn', 'aws', 'gcp', 'sdk', 'ide', 'cli', 'gui',
        'ctr', 'kpi', 'roi', 'etl', 'crm', 'erp', 'cms', 'iot', 'nlp', 'ml', 'ai', 'ml',
        'devops', 'cicd', 'agile', 'scrum', 'kanban', 'mvp', 'poc', 'qa', 'qc', 'sqa',
        'backend', 'frontend', 'fullstack', 'saas', 'paas', 'iaas', 'nosql', 'rdbms'
    }
    
    # Immediate rejection for technical terms
    if text_lower in technical_terms:
        return False
    
    # Strong exclusions - definitive non-names
    strong_exclusions = [
        # Resume sections
        'objective', 'summary', 'about', 'skills', 'education', 'experience', 'projects', 'resume', 'cv',
        'certifications', 'achievements', 'interests', 'languages', 'profile', 'career', 'personal',
        'responsibilities', 'duties', 'accomplishments', 'awards', 'honors', 'activities',
        
        # Academic terms
        'coursework', 'relevant', 'course', 'courses', 'curriculum', 'academic', 'semester', 'year',
        'gpa', 'grade', 'marks', 'score', 'cgpa', 'percentage', 'transcript', 'syllabus', 'subject',
        'major', 'minor', 'specialization', 'concentration', 'thesis', 'dissertation', 'research',
        'internship', 'training', 'workshop', 'seminar', 'conference', 'publication', 'paper',
        
        # Professional terms
        'position', 'designation', 'title', 'department', 'company', 'organization', 'employer',
        'manager', 'supervisor', 'team', 'project', 'developer', 'engineer', 'analyst', 'consultant',
        
        # Contact terms
        'phone', 'email', 'mobile', 'contact', 'linkedin', 'github', 'website', 'www', 'http', 'tel',
        'address', 'street', 'road', 'avenue', 'lane', 'city', 'state', 'country', 'pin', 'zip',
        
        # Technical terms
        'programming', 'software', 'hardware', 'database', 'network', 'system', 'application',
        'development', 'testing', 'debugging', 'coding', 'scripting', 'framework', 'library',
        
        # Common false positives
        'india', 'indian', 'engineering', 'college', 'university', 'institute', 'technology', 'science',
        'bachelor', 'master', 'degree', 'diploma', 'certificate', 'graduation'
    ]
    
    # Check for strong exclusions
    if any(exclusion in text_lower for exclusion in strong_exclusions):
        return False
    
    # Specific phrase exclusions
    excluded_phrases = [
        'relevant coursework', 'course work', 'academic coursework', 'professional experience',
        'work experience', 'technical skills', 'soft skills', 'computer skills', 'programming skills',
        'project management', 'team leadership', 'problem solving', 'data analysis', 'machine learning',
        'artificial intelligence', 'web development', 'software development', 'database management',
        'network administration', 'quality assurance', 'user interface', 'user experience',
        'business analysis', 'system administration', 'cloud computing', 'cyber security'
    ]
    
    if text_lower in excluded_phrases:
        return False
    
    # Additional checks for company names
    if (len(words) > 2 and  # Companies often have 3+ words
        any(word.isupper() and len(word) > 1 for word in words) and  # All caps words
        not all(word[0].isupper() and word[1:].islower() for word in words if word.isalpha())):  # Not proper name format
        return False
    
    # Pattern-based validation for actual names
    name_patterns = [
        r'^[A-Z][a-z]+(?:\s[A-Z][a-z]+){0,3}$',  # Title case names (John Smith)
        r'^[A-Z][a-z]+(?:\s[A-Z]\.){0,2}(?:\s[A-Z][a-z]+)?$'  # Names with middle initials (John A. Smith)
    ]
    
    # Must match at least one name pattern
    if not any(re.match(pattern, text_clean) for pattern in name_patterns):
        return False
    
    # Validate each word
    for word in words:
        if not word or len(word) < 2 or len(word) > 20:
            return False
        
        # Must be alphabetic or single letter followed by dot (middle initial)
        if not (word.isalpha() or (len(word) == 2 and word[1] == '.' and word[0].isalpha())):
            return False
        
        # Check if word looks like a common name part
        word_lower = word.lower().rstrip('.')
        
        # Exclude technical terms at word level too
        if word_lower in technical_terms or word_lower in company_indicators:
            return False
        
        # Exclude obvious non-name words
        non_name_words = {
            'the', 'and', 'or', 'but', 'for', 'with', 'from', 'into', 'upon', 'over', 'under',
            'above', 'below', 'through', 'during', 'before', 'after', 'between', 'among',
            'this', 'that', 'these', 'those', 'what', 'which', 'who', 'where', 'when', 'why',
            'how', 'all', 'any', 'each', 'every', 'some', 'many', 'much', 'few', 'little',
            'more', 'most', 'less', 'least', 'good', 'better', 'best', 'bad', 'worse', 'worst'
        }
        
        if word_lower in non_name_words:
            return False
    
    return True

def extract_name(text, entities):
    """Ultra-accurate name extraction with comprehensive validation"""
    lines = [line.strip() for line in text.splitlines()[:15] if line.strip()]
    
    # Debug: Print the first few lines to understand the structure
    print(f"DEBUG - Resume first 5 lines:")
    for i, line in enumerate(lines[:5]):
        print(f"  Line {i}: '{line}'")
    
    # Known location/place names to exclude (common false positives)
    location_names = {
        'puducherry', 'chennai', 'mumbai', 'delhi', 'bangalore', 'hyderabad', 'kolkata', 'pune', 'ahmedabad',
        'surat', 'jaipur', 'lucknow', 'kanpur', 'nagpur', 'indore', 'thane', 'bhopal', 'visakhapatnam',
        'pimpri', 'patna', 'vadodara', 'ghaziabad', 'ludhiana', 'agra', 'nashik', 'faridabad', 'meerut',
        'rajkot', 'kalyan', 'vasai', 'varanasi', 'srinagar', 'aurangabad', 'dhanbad', 'amritsar', 'navi',
        'allahabad', 'ranchi', 'howrah', 'coimbatore', 'jabalpur', 'gwalior', 'vijayawada', 'jodhpur',
        'madurai', 'raipur', 'kota', 'guwahati', 'chandigarh', 'solapur', 'hubballi', 'tiruchirappalli',
        'bareilly', 'mysore', 'tiruppur', 'gurgaon', 'aligarh', 'jalandhar', 'bhubaneswar', 'salem',
        'warangal', 'guntur', 'bhiwandi', 'saharanpur', 'gorakhpur', 'bikaner', 'amravati', 'noida',
        'jamshedpur', 'bhilai', 'cuttack', 'firozabad', 'kochi', 'nellore', 'bhavnagar', 'dehradun',
        'durgapur', 'asansol', 'rourkela', 'nanded', 'kolhapur', 'ajmer', 'akola', 'gulbarga', 'jamnagar',
        'ujjain', 'loni', 'siliguri', 'jhansi', 'ulhasnagar', 'jammu', 'sangli', 'mangalore', 'erode',
        'belgaum', 'ambattur', 'tirunelveli', 'malegaon', 'gaya', 'jalgaon', 'udaipur', 'maheshtala',
        'india', 'united states', 'usa', 'uk', 'canada', 'australia', 'singapore', 'malaysia', 'thailand'
    }
    
    # Company names and organization patterns to exclude
    company_exclusions = {
        'codtech', 'it', 'solutions', 'technologies', 'systems', 'services', 'consulting',
        'software', 'tech', 'corporation', 'corp', 'inc', 'ltd', 'limited', 'pvt', 'private',
        'company', 'enterprises', 'group', 'international', 'global', 'industries', 'infosys',
        'tcs', 'wipro', 'accenture', 'cognizant', 'capgemini', 'deloitte', 'microsoft', 'google',
        'amazon', 'apple', 'facebook', 'meta', 'netflix', 'oracle', 'ibm', 'adobe', 'salesforce'
    }
    
    # Strategy 1: Simple pattern matching for names in first few lines
    def looks_like_name(text):
        """Simple check if text looks like a personal name"""
        if not text or len(text.strip()) < 2:
            return False
        
        words = text.strip().split()
        if len(words) < 1 or len(words) > 4:
            return False
        
        # Check if it contains numbers, symbols, or common non-name patterns
        if any(char.isdigit() or char in '@#$%&*()[]{}' for char in text):
            return False
        
        # Check for company indicators
        text_lower = text.lower()
        if any(comp in text_lower for comp in company_exclusions):
            return False
        
        # Check if it's a location
        if text_lower in location_names:
            return False
        
        # Simple name pattern: starts with capital letters
        for word in words:
            if not word[0].isupper():
                return False
        
        return True
    
    # Check first few lines for name patterns
    for i, line in enumerate(lines[:5]):
        line_clean = line.strip()
        
        # Skip empty lines
        if not line_clean:
            continue
        
        # Skip lines that are clearly headers or contact info
        line_lower = line_clean.lower()
        if (any(keyword in line_lower for keyword in ['resume', 'cv', 'phone', 'email', '@', 'linkedin', 'github', 'http']) or
            len([c for c in line_clean if c.isdigit()]) > 3):  # Too many numbers
            continue
        
        # Check if the entire line could be a name
        if looks_like_name(line_clean) and len(line_clean.split()) <= 3:
            print(f"DEBUG - Found potential name in line {i}: '{line_clean}'")
            return line_clean
        
        # Check if the line starts with a name (extract first 2-3 words)
        words = line_clean.split()
        if len(words) >= 2:
            potential_name = ' '.join(words[:3])  # Take first 3 words maximum
            if looks_like_name(potential_name):
                print(f"DEBUG - Found potential name at start of line {i}: '{potential_name}'")
                return potential_name
        
        # Single word names (first names only)
        if len(words) >= 1 and looks_like_name(words[0]) and len(words[0]) > 2:
            # Check if it's followed by a last name or initial
            if len(words) > 1 and (looks_like_name(words[1]) or (len(words[1]) == 1 and words[1].isupper())):
                name_candidate = f"{words[0]} {words[1]}"
                print(f"DEBUG - Found name candidate: '{name_candidate}'")
                return name_candidate
    
    # Strategy 2: Look for PERSON entities but with more lenient validation
    person_entities = [(ent_text, ent_label) for ent_text, ent_label in entities if ent_label == "PERSON"]
    
    for ent_text, ent_label in person_entities:
        # Basic validation
        if (ent_text and 
            len(ent_text.split()) <= 4 and
            not any(char.isdigit() for char in ent_text) and
            ent_text.lower() not in location_names and
            not any(comp in ent_text.lower() for comp in company_exclusions)):
            
            print(f"DEBUG - Found valid PERSON entity: '{ent_text}'")
            return ent_text
    
    print("DEBUG - No valid name found, returning 'Not found'")
    return "Not found"

def extract_location(text, entities):
    """Enhanced location extraction with multiple strategies"""
    
    # Strategy 1: Try spaCy GPE entities first with better filtering
    for ent_text, ent_label in entities:
        if (ent_label == "GPE" and 
            ent_text.strip().upper() not in ["NO", "YES", "LINKEDIN", "GITHUB", "FACEBOOK", "RESUME", "CV"] and
            ent_text.upper() not in IGNORE_LIST and
            len(ent_text.strip()) > 2 and
            not ent_text.isdigit()):
            return ent_text
    
    # Strategy 2: Look for explicit address/location patterns
    location_patterns = [
        r"Address[:\-\s]*(.+?)(?:\n|$)",
        r"Location[:\-\s]*(.+?)(?:\n|$)", 
        r"City[:\-\s]*(.+?)(?:\n|$)",
        r"([A-Za-z\s]+,\s*[A-Za-z\s]+(?:,\s*\d{5,6})?)",  # City, State format
        r"([A-Za-z\s]+\s+\d{5,6})",  # City with PIN/ZIP
        r"(\w+(?:\s+\w+)*,\s*\w+(?:\s+\w+)*)"  # General comma-separated format
    ]
    
    for pattern in location_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            location = match.group(1).strip()
            # Clean up and validate location
            if (len(location) > 3 and 
                not location.lower().startswith(('phone', 'email', 'mobile', 'tel')) and
                not re.match(r'^[\d\-\+\(\)\s]+$', location)):  # Not just numbers/phone
                return location
    
    # Strategy 3: Look for lines with address/location keywords
    for line in text.splitlines():
        line_lower = line.lower().strip()
        if any(word in line_lower for word in ["address", "location", "city", "state", "district", "pincode", "pin"]):
            # Extract location after the keyword
            parts = re.split(r'[:\-]', line, 1)
            if len(parts) > 1:
                location = parts[1].strip()
                if (len(location) > 2 and 
                    not re.match(r'^[\d\-\+\(\)\s]+$', location) and  # Not just phone number
                    not any(keyword in location.lower() for keyword in ['phone', 'email', 'mobile'])):
                    return location
            else:
                # Remove common prefixes and return clean location
                clean_line = re.sub(r'^(address|location|city|state|district|pincode|pin)[:\-\s]*', '', line, flags=re.IGNORECASE).strip()
                if (len(clean_line) > 2 and 
                    not re.match(r'^[\d\-\+\(\)\s]+$', clean_line)):
                    return clean_line
    
    # Strategy 4: Look for location patterns in any line
    for line in text.splitlines()[:10]:  # Check first 10 lines
        line = line.strip()
        # Match common location formats like "City, State" or "City - State"
        location_match = re.search(r'^([A-Za-z\s]+(?:[,\-]\s*[A-Za-z\s]+)+)', line)
        if location_match:
            location = location_match.group(1).strip()
            if (len(location) > 5 and 
                not any(keyword in location.lower() for keyword in ['phone', 'email', 'mobile', 'resume', 'cv', 'linkedin', 'github'])):
                return location
    
    return "Not found"

def extract_dates(entities):
    """Extract meaningful dates (graduation years, work periods) from entities"""
    import datetime
    current_year = datetime.datetime.now().year
    dates = []
    seen_dates = set()
    
    for text, label in entities:
        if label == "DATE":
            # Look for graduation years (typically 4 years ago to present)
            year_match = re.search(r'\b(20\d{2})\b', text)
            if year_match:
                year = int(year_match.group(1))
                # Only include reasonable years (graduation/work years)
                if 2015 <= year <= current_year and year not in seen_dates:
                    dates.append(str(year))
                    seen_dates.add(year)
    
    # Limit to most recent 3 years to avoid clutter
    dates = sorted(dates, reverse=True)[:3]
    return ", ".join(dates) if dates else "Not found"

def extract_percent(entities):
    for text, label in entities:
        if label == "PERCENT":
            return text
    return "Not found"

def generate_summary(entities, resume_text):
    """Generate comprehensive resume summary"""
    name = extract_name(resume_text, entities)
    location = extract_location(resume_text, entities)
    dates = extract_dates(entities)
    percent = extract_percent(entities)
    
    # Get additional analysis
    content_analysis = analyze_resume_content(resume_text)
    
    # Debug print to verify correct values
    print(f"DEBUG - Resume text (first 500 chars): {resume_text[:500]}")
    print(f"DEBUG - Education Level: {content_analysis['education_level']}")
    print(f"DEBUG - Experience Level: {content_analysis['experience_level']}")
    
    return {
        "name": name,
        "locations": location,
        "dates": dates,
        "percent": percent,
        "skills": content_analysis['skills'],
        "education_level": content_analysis['education_level'],
        "experience_level": content_analysis['experience_level']
    }

def analyze_entities(doc, resume_text):
    """Analyze entities and predict job role using ML"""
    entities = [(ent.text, ent.label_) for ent in doc.ents if ent.label_ != 'ORG']
    
    # Get ML-based job prediction
    predicted_role, ml_confidence = predict_job_role(resume_text)
    
    # Enhanced scoring with more realistic weights and bonuses
    completeness_factors = {
        'name': 15,      # Name found (reduced weight)
        'location': 10,  # Location found (reduced weight) 
        'dates': 10,     # Education/work dates (reduced weight)
        'skills': 30,    # Technical skills found (increased weight)
        'experience': 20, # Experience level detected (increased weight)
        'achievements': 15 # Percentages/achievements (increased weight)
    }
    
    # Base score starts higher for any valid resume
    base_score = 35  # Every resume starts with 35% base score
    
    # Initialize score breakdown for transparency
    score_breakdown = {
        'ml_confidence': ml_confidence,
        'completeness_score': 0,
        'base_score': base_score,
        'components': {
            'name': {'score': 0, 'max': 15, 'found': False, 'value': "Not found"},
            'location': {'score': 0, 'max': 10, 'found': False, 'value': "Not found"},
            'dates': {'score': 0, 'max': 10, 'found': False, 'value': "Not found"},
            'skills': {'score': 0, 'max': 30, 'found': False, 'value': []},
            'experience': {'score': 0, 'max': 20, 'found': False, 'value': "Entry Level"},
            'achievements': {'score': 0, 'max': 15, 'found': False, 'value': "Not found"}
        }
    }
    
    completeness_score = base_score  # Start with base score
    content_analysis = analyze_resume_content(resume_text)
    
    # Check each factor and update breakdown with enhanced scoring
    name_value = extract_name(resume_text, entities)
    if name_value != "Not found":
        completeness_score += completeness_factors['name']
        score_breakdown['components']['name']['score'] = completeness_factors['name']
        score_breakdown['components']['name']['found'] = True
        score_breakdown['components']['name']['value'] = name_value
    else:
        # Partial credit if resume has some structure
        if len(resume_text.split()) > 50:  # Has substantial content
            completeness_score += completeness_factors['name'] * 0.3  # 30% partial credit
            score_breakdown['components']['name']['score'] = int(completeness_factors['name'] * 0.3)

    location_value = extract_location(resume_text, entities)
    if location_value != "Not found":
        completeness_score += completeness_factors['location']
        score_breakdown['components']['location']['score'] = completeness_factors['location']
        score_breakdown['components']['location']['found'] = True
        score_breakdown['components']['location']['value'] = location_value

    dates_value = extract_dates(entities)
    if dates_value != "Not found":
        completeness_score += completeness_factors['dates']
        score_breakdown['components']['dates']['score'] = completeness_factors['dates']
        score_breakdown['components']['dates']['found'] = True
        score_breakdown['components']['dates']['value'] = dates_value
    
    # Enhanced skills scoring with progressive bonuses
    if content_analysis['skills']:
        skill_count = len(content_analysis['skills'])
        if skill_count >= 5:
            skill_score = completeness_factors['skills']  # Full points for 5+ skills
        elif skill_count >= 3:
            skill_score = int(completeness_factors['skills'] * 0.8)  # 80% for 3-4 skills
        else:
            skill_score = int(completeness_factors['skills'] * 0.5)  # 50% for 1-2 skills
        
        completeness_score += skill_score
        score_breakdown['components']['skills']['score'] = skill_score
        score_breakdown['components']['skills']['found'] = True
        score_breakdown['components']['skills']['value'] = content_analysis['skills']
    else:
        # Check if resume mentions any technical terms for partial credit
        tech_indicators = ['python', 'java', 'programming', 'software', 'development', 'coding', 
                          'database', 'web', 'mobile', 'application', 'system', 'technology']
        resume_lower = resume_text.lower()
        if any(indicator in resume_lower for indicator in tech_indicators):
            partial_skill_score = int(completeness_factors['skills'] * 0.3)
            completeness_score += partial_skill_score
            score_breakdown['components']['skills']['score'] = partial_skill_score
    
    # Enhanced experience scoring
    experience_level = content_analysis['experience_level']
    if experience_level == "Senior Level":
        experience_score = completeness_factors['experience']
        completeness_score += experience_score
        score_breakdown['components']['experience']['score'] = experience_score
        score_breakdown['components']['experience']['found'] = True
    elif experience_level == "Mid Level":
        experience_score = int(completeness_factors['experience'] * 0.8)
        completeness_score += experience_score
        score_breakdown['components']['experience']['score'] = experience_score
        score_breakdown['components']['experience']['found'] = True
    else:  # Entry Level
        experience_score = int(completeness_factors['experience'] * 0.5)
        completeness_score += experience_score
        score_breakdown['components']['experience']['score'] = experience_score
    
    score_breakdown['components']['experience']['value'] = experience_level
    
    achievements_value = extract_percent(entities)
    if achievements_value != "Not found":
        completeness_score += completeness_factors['achievements']
        score_breakdown['components']['achievements']['score'] = completeness_factors['achievements']
        score_breakdown['components']['achievements']['found'] = True
        score_breakdown['components']['achievements']['value'] = achievements_value
    
    # Content quality bonuses
    word_count = len(resume_text.split())
    if word_count > 200:  # Substantial resume content
        completeness_score += 5
    if word_count > 400:  # Very detailed resume
        completeness_score += 5
    
    # Education level bonus
    edu_level = content_analysis['education_level']
    if edu_level != "Not specified":
        if "bachelor" in edu_level.lower() or "master" in edu_level.lower() or "phd" in edu_level.lower():
            completeness_score += 5
        else:
            completeness_score += 2
    
    # Cap the completeness score at 100
    completeness_score = min(completeness_score, 100)
    score_breakdown['completeness_score'] = completeness_score
    
    # Improved final scoring algorithm
    # Increase ML weight for better job matching, reduce completeness dependency
    enhanced_ml_confidence = min(ml_confidence * 1.2, 100)  # Boost ML confidence slightly
    
    # New weighted formula: 70% Enhanced ML + 30% Completeness
    match_percentage = int((enhanced_ml_confidence * 0.7) + (completeness_score * 0.3))
    
    # Ensure reasonable minimum scores based on content quality
    if completeness_score >= 80:
        match_percentage = max(match_percentage, 70)  # High quality resumes get at least 70%
    elif completeness_score >= 60:
        match_percentage = max(match_percentage, 60)  # Good resumes get at least 60%
    elif completeness_score >= 40:
        match_percentage = max(match_percentage, 50)  # Basic resumes get at least 50%
    else:
        match_percentage = max(match_percentage, 35)  # Even poor resumes get something
    
    # Cap final score at 100
    match_percentage = min(match_percentage, 100)
    
    
    score_breakdown['final_score'] = match_percentage
    score_breakdown['ml_weight'] = 0.7
    score_breakdown['completeness_weight'] = 0.3
    score_breakdown['completeness_percentage'] = completeness_score
    score_breakdown['enhanced_ml_confidence'] = enhanced_ml_confidence
    
    # Generate suggestions based on missing information
    missing_info = []
    if extract_name(resume_text, entities) == "Not found":
        missing_info.append("NAME")
    if extract_location(resume_text, entities) == "Not found":
        missing_info.append("LOCATION")
    if extract_dates(entities) == "Not found":
        missing_info.append("DATES")
    if not content_analysis['skills']:
        missing_info.append("TECHNICAL_SKILLS")
    if extract_percent(entities) == "Not found":
        missing_info.append("ACHIEVEMENTS")
    
    return entities, match_percentage, missing_info, predicted_role, score_breakdown

def generate_keyword_suggestions(predicted_role, content_analysis, resume_text):
    """Generate relevant keywords and suggestions based on job role and current resume content"""
    
    # Job-specific keyword databases
    job_keywords = {
        "Software Engineer": {
            "technical_skills": [
                "Python", "Java", "JavaScript", "C++", "React", "Node.js", "SQL", "Git", 
                "Docker", "Kubernetes", "AWS", "CI/CD", "Agile", "Scrum", "REST API", 
                "Microservices", "TDD", "Object-Oriented Programming", "Data Structures", "Algorithms"
            ],
            "soft_skills": [
                "Problem-solving", "Team collaboration", "Communication", "Leadership", 
                "Critical thinking", "Adaptability", "Time management", "Project management"
            ],
            "action_verbs": [
                "Developed", "Implemented", "Designed", "Optimized", "Built", "Created", 
                "Maintained", "Integrated", "Deployed", "Collaborated", "Led", "Managed"
            ],
            "certifications": [
                "AWS Certified Developer", "Oracle Java Certification", "Microsoft Azure Fundamentals",
                "Google Cloud Professional", "Scrum Master Certification"
            ]
        },
        "Frontend Developer": {
            "technical_skills": [
                "HTML5", "CSS3", "JavaScript", "React", "Angular", "Vue.js", "TypeScript", 
                "SASS/SCSS", "Webpack", "Bootstrap", "Responsive Design", "Cross-browser Compatibility",
                "Performance Optimization", "Progressive Web Apps", "Material-UI", "Tailwind CSS"
            ],
            "soft_skills": [
                "UI/UX Design", "Attention to detail", "Creativity", "User-centered thinking",
                "Cross-functional collaboration", "Communication", "Problem-solving"
            ],
            "action_verbs": [
                "Designed", "Developed", "Implemented", "Optimized", "Created", "Built", 
                "Enhanced", "Improved", "Collaborated", "Delivered"
            ],
            "certifications": [
                "Google UX Design Certificate", "Adobe Certified Expert", "W3C Frontend Certification",
                "React Developer Certification"
            ]
        },
        "Backend Developer": {
            "technical_skills": [
                "Python", "Java", "Node.js", "C#", ".NET", "SQL", "NoSQL", "MongoDB", "PostgreSQL",
                "REST API", "GraphQL", "Microservices", "Docker", "Kubernetes", "AWS", "Azure",
                "Redis", "Message Queues", "Database Design", "System Architecture"
            ],
            "soft_skills": [
                "System thinking", "Problem-solving", "Performance optimization", "Security awareness",
                "Team collaboration", "Documentation", "Code review"
            ],
            "action_verbs": [
                "Architected", "Developed", "Implemented", "Optimized", "Designed", "Built",
                "Maintained", "Scaled", "Integrated", "Deployed"
            ],
            "certifications": [
                "AWS Solutions Architect", "Microsoft Azure Developer", "Oracle Database Certification",
                "MongoDB Developer Certification"
            ]
        },
        "Data Scientist": {
            "technical_skills": [
                "Python", "R", "SQL", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch",
                "Pandas", "NumPy", "Scikit-learn", "Matplotlib", "Seaborn", "Jupyter", "Statistics",
                "Data Visualization", "Big Data", "Hadoop", "Spark", "Tableau", "Power BI"
            ],
            "soft_skills": [
                "Analytical thinking", "Statistical analysis", "Business acumen", "Communication",
                "Data storytelling", "Problem-solving", "Research methodology"
            ],
            "action_verbs": [
                "Analyzed", "Modeled", "Predicted", "Optimized", "Visualized", "Researched",
                "Implemented", "Built", "Developed", "Improved"
            ],
            "certifications": [
                "Google Data Analytics Certificate", "IBM Data Science Certification", 
                "Microsoft Azure Data Scientist", "Tableau Certification"
            ]
        },
        "DevOps Engineer": {
            "technical_skills": [
                "Docker", "Kubernetes", "Jenkins", "CI/CD", "AWS", "Azure", "GCP", "Terraform",
                "Ansible", "Linux", "Shell Scripting", "Python", "Monitoring", "Prometheus",
                "Grafana", "Infrastructure as Code", "GitOps", "Microservices"
            ],
            "soft_skills": [
                "Process improvement", "Automation mindset", "Troubleshooting", "Collaboration",
                "System thinking", "Continuous learning", "Problem-solving"
            ],
            "action_verbs": [
                "Automated", "Deployed", "Configured", "Optimized", "Monitored", "Implemented",
                "Maintained", "Scaled", "Integrated", "Streamlined"
            ],
            "certifications": [
                "AWS DevOps Engineer", "Azure DevOps Engineer", "Docker Certified Associate",
                "Kubernetes Administrator Certification"
            ]
        }
    }
    
    # Default keywords for general roles
    default_keywords = {
        "technical_skills": [
            "Programming", "Software Development", "Problem Solving", "Testing", "Debugging",
            "Version Control", "Documentation", "Code Review", "Agile Methodology"
        ],
        "soft_skills": [
            "Communication", "Teamwork", "Leadership", "Problem-solving", "Adaptability",
            "Time management", "Critical thinking", "Collaboration"
        ],
        "action_verbs": [
            "Developed", "Implemented", "Created", "Managed", "Led", "Collaborated",
            "Improved", "Optimized", "Delivered", "Achieved"
        ],
        "certifications": [
            "Industry-relevant certifications", "Professional development courses",
            "Technical training programs"
        ]
    }
    
    # Get keywords for the predicted role
    role_keywords = job_keywords.get(predicted_role, default_keywords)
    
    # Analyze current resume content
    resume_lower = resume_text.lower()
    current_skills = [skill.lower() for skill in content_analysis['skills']]
    
    # Find missing technical keywords
    missing_technical = []
    for keyword in role_keywords["technical_skills"]:
        if keyword.lower() not in resume_lower and keyword.lower() not in current_skills:
            missing_technical.append(keyword)
    
    # Find missing soft skills
    missing_soft = []
    for keyword in role_keywords["soft_skills"]:
        if keyword.lower() not in resume_lower:
            missing_soft.append(keyword)
    
    # Find missing action verbs
    missing_verbs = []
    for verb in role_keywords["action_verbs"]:
        if verb.lower() not in resume_lower:
            missing_verbs.append(verb)
    
    # Generate improvement suggestions
    improvement_suggestions = generate_improvement_descriptions(
        predicted_role, content_analysis, missing_technical, missing_soft, missing_verbs
    )
    
    return {
        "role": predicted_role,
        "missing_technical_skills": missing_technical[:8],  # Top 8 suggestions
        "missing_soft_skills": missing_soft[:6],           # Top 6 suggestions
        "missing_action_verbs": missing_verbs[:6],         # Top 6 suggestions
        "suggested_certifications": role_keywords["certifications"][:4],  # Top 4 suggestions
        "improvement_suggestions": improvement_suggestions,
        "current_skills_count": len(content_analysis['skills']),
        "skill_gap_analysis": {
            "strong_areas": content_analysis['skills'],
            "improvement_areas": missing_technical[:5]
        }
    }

def generate_improvement_descriptions(predicted_role, content_analysis, missing_technical, missing_soft, missing_verbs):
    """Generate detailed descriptions on how to improve resume score"""
    
    suggestions = []
    
    # Technical Skills Improvement
    if len(content_analysis['skills']) < 5:
        suggestions.append({
            "category": "Technical Skills",
            "priority": "High",
            "description": f"Add more technical skills relevant to {predicted_role}. You currently have {len(content_analysis['skills'])} skills listed. Aim for 8-12 key technical skills.",
            "action_items": [
                f"Include programming languages: {', '.join(missing_technical[:3])}",
                "Add frameworks and tools you've used",
                "Mention databases and cloud platforms",
                "Include development methodologies (Agile, Scrum)"
            ],
            "impact": "Can increase your score by up to 15 points"
        })
    
    # Experience Enhancement
    if content_analysis['experience_level'] == "Entry Level":
        suggestions.append({
            "category": "Experience Description",
            "priority": "High",
            "description": "Enhance your experience descriptions with quantifiable achievements and impact metrics.",
            "action_items": [
                "Add specific numbers and percentages to achievements",
                "Use strong action verbs to start bullet points",
                "Describe the impact of your work on business/team",
                "Include technologies used in each role"
            ],
            "impact": "Can increase your score by up to 10 points"
        })
    
    # Action Verbs Improvement
    if len(missing_verbs) > 3:
        suggestions.append({
            "category": "Action Verbs",
            "priority": "Medium",
            "description": "Use stronger action verbs to make your resume more impactful and dynamic.",
            "action_items": [
                f"Replace weak verbs with: {', '.join(missing_verbs[:4])}",
                "Start each bullet point with a strong action verb",
                "Avoid repetitive verbs throughout your resume",
                "Use past tense for previous roles, present for current"
            ],
            "impact": "Improves overall resume readability and impact"
        })
    
    # Soft Skills Enhancement
    if len(missing_soft) > 3:
        suggestions.append({
            "category": "Soft Skills Integration",
            "priority": "Medium",
            "description": "Integrate soft skills naturally into your experience descriptions rather than listing them separately.",
            "action_items": [
                f"Demonstrate skills like: {', '.join(missing_soft[:3])}",
                "Show leadership through project examples",
                "Highlight collaboration in team settings",
                "Mention communication in client/stakeholder interactions"
            ],
            "impact": "Makes your profile more well-rounded"
        })
    
    # Keywords Optimization
    suggestions.append({
        "category": "Keyword Optimization",
        "priority": "High",
        "description": f"Optimize your resume with {predicted_role}-specific keywords to improve ATS compatibility.",
        "action_items": [
            "Include industry-standard terminology",
            "Match keywords from job descriptions you're targeting",
            "Use both acronyms and full forms (e.g., 'AI' and 'Artificial Intelligence')",
            "Naturally integrate keywords into context"
        ],
        "impact": "Significantly improves ATS scanning and recruiter matching"
    })
    
    # Quantifiable Achievements
    suggestions.append({
        "category": "Quantifiable Results",
        "priority": "High",
        "description": "Add more numbers, percentages, and measurable outcomes to demonstrate your impact.",
        "action_items": [
            "Include performance improvements (e.g., '30% faster processing')",
            "Mention team sizes you've led or worked with",
            "Add project budgets, timelines, or user counts",
            "Quantify cost savings or revenue increases"
        ],
        "impact": "Can increase your score by up to 10 points"
    })
    
    return suggestions

def get_suggestions(missing_entities, job_role):
    """Generate actionable suggestions based on missing information and predicted job role"""
    suggestions = []
    
    # General suggestions based on missing information
    for ent in missing_entities:
        if ent == "NAME":
            suggestions.append("Add your full name prominently at the top of your resume.")
        elif ent == "DATES":
            suggestions.append("Include graduation year and work experience dates to show career progression.")
        elif ent == "LOCATION":
            suggestions.append("Mention your current location or preferred work location.")
        elif ent == "ACHIEVEMENTS":
            suggestions.append("Add quantifiable achievements with percentages, numbers, or metrics.")
        elif ent == "TECHNICAL_SKILLS":
            suggestions.append("Include a dedicated skills section with relevant technical skills.")
    
    # Job-specific suggestions
    job_specific_tips = {
        "Data Scientist": [
            "Highlight your experience with Python, R, SQL, and machine learning frameworks.",
            "Include specific projects showing data analysis and statistical modeling skills."
        ],
        "Software Engineer": [
            "Emphasize your programming languages and software development experience.",
            "Include details about software projects and development methodologies."
        ],
        "Frontend Developer": [
            "Highlight your HTML, CSS, JavaScript skills and modern frameworks like React or Angular.",
            "Include portfolio links or examples of responsive web design projects."
        ],
        "Backend Developer": [
            "Emphasize server-side programming languages and database management skills.",
            "Include experience with API development and system architecture."
        ],
        "DevOps Engineer": [
            "Highlight experience with containerization (Docker, Kubernetes) and cloud platforms.",
            "Include details about CI/CD pipelines and infrastructure automation."
        ]
    }
    
    # Add job-specific suggestions
    if job_role in job_specific_tips:
        suggestions.extend(job_specific_tips[job_role][:2])
    
    # General improvement suggestions
    if not missing_entities:
        suggestions.extend([
            "Consider adding more technical skills relevant to your target role.",
            "Include quantifiable achievements and impact metrics in your experience section."
        ])
    
    return suggestions[:5]

@app.route('/', methods=['GET', 'POST'])
def index():
    entities = []
    resume_text = ""
    match_percentage = 0
    missing_entities = []
    job_match = ""
    summary = {}
    suggestions = []
    nltk_analysis = {}
    score_breakdown = {}
    keyword_suggestions = {}
    
    if request.method == 'POST':
        uploaded_file = request.files.get('resume_file')
        manual_name = request.form.get('manual_name', '').strip()
        manual_location = request.form.get('manual_location', '').strip()
        manual_dates = request.form.get('manual_dates', '').strip()
        manual_percent = request.form.get('manual_percent', '').strip()
        
        if uploaded_file and uploaded_file.filename:
            resume_text = extract_text(uploaded_file)
            
            # Use appropriate NLP system (spaCy or fallback)
            doc = nlp(resume_text)
            
            # Use ML-based analysis
            entities, match_percentage, missing_entities, job_match, score_breakdown = analyze_entities(doc, resume_text)
            summary = generate_summary(entities, resume_text)
            
            # Generate keyword suggestions and improvement guidance
            content_analysis = analyze_resume_content(resume_text)
            keyword_suggestions = generate_keyword_suggestions(job_match, content_analysis, resume_text)
            
            # Enhanced NLTK analysis (with fallback)
            try:
                nltk_analysis = analyze_resume_with_nltk(resume_text)
            except Exception as e:
                print(f"NLTK analysis failed: {e}")
                # Provide fallback analysis without NLTK
                nltk_analysis = {
                    'sentiment': {'positive': 0.5, 'negative': 0.1, 'neutral': 0.4, 'compound': 0.3},
                    'entities': [],
                    'statistics': {
                        'word_count': len(resume_text.split()),
                        'sentence_count': len(resume_text.split('.')),
                        'avg_sentence_length': len(resume_text.split()) / max(len(resume_text.split('.')), 1)
                    },
                    'key_phrases': [],
                    'filtered_words': []
                }
            
            # Manual corrections
            if manual_name:
                summary['name'] = manual_name
            if manual_location:
                summary['locations'] = manual_location
            if manual_dates:
                summary['dates'] = manual_dates
            if manual_percent:
                summary['percent'] = manual_percent
            
            suggestions = get_suggestions(missing_entities, job_match)
    
    return render_template(
        'index.html',
        entities=entities,
        resume_text=resume_text,
        match_percentage=match_percentage,
        missing_entities=missing_entities,
        job_match=job_match,
        summary=summary,
        suggestions=suggestions,
        nltk_analysis=nltk_analysis,
        score_breakdown=score_breakdown,
        keyword_suggestions=keyword_suggestions
    )

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    host = '0.0.0.0'
    debug = os.environ.get('FLASK_ENV') != 'production'
    
    print(f"Starting Flask app on {host}:{port} (debug={debug})")
    app.run(host=host, port=port, debug=debug)
